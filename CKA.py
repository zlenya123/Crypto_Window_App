from PyQt5 import QtWidgets, QtCore
from PyQt5.QtGui import QIcon
from collections import Counter
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
import docx
import os

RUSSIAN_PROBABILITIES = {
    'о': 0.090, 'е': 0.072, 'а': 0.062, 'и': 0.062, 'н': 0.053,
    'т': 0.053, 'с': 0.045, 'р': 0.040, 'в': 0.038, 'л': 0.035,
    'к': 0.028, 'м': 0.026, 'д': 0.025, 'п': 0.023, 'у': 0.021,
    'я': 0.018, 'ы': 0.016, 'з': 0.016, 'ь': 0.014, 'б': 0.014,
    'г': 0.013, 'ч': 0.012, 'й': 0.010, 'х': 0.009, 'ж': 0.007,
    'ю': 0.006, 'ш': 0.006, 'ц': 0.004, 'щ': 0.003, 'э': 0.003,
    'ф': 0.002, 'ё': 0.001, 'ъ': 0.001 
}

ENGLISH_PROBABILITIES = {
    'e': 0.123, 't': 0.096, 'a': 0.081, 'o': 0.079, 'n': 0.072,
    'i': 0.071, 's': 0.066, 'r': 0.060, 'h': 0.051, 'l': 0.040,
    'd': 0.036, 'c': 0.032, 'u': 0.031, 'p': 0.023, 'f': 0.023,
    'm': 0.022, 'w': 0.020, 'y': 0.019, 'b': 0.016, 'g': 0.016,
    'v': 0.009, 'k': 0.005, 'x': 0.002, 'q': 0.002, 'j': 0.001,
    'z': 0.001
}

class HistogramWidget(QtWidgets.QDialog):
    def __init__(self, freq_data, title):
        super().__init__()
        self.setWindowModality(QtCore.Qt.WindowModal)
        self.figure = Figure(figsize=(8, 5))
        self.canvas = FigureCanvas(self.figure)
        layout = QtWidgets.QVBoxLayout()
        layout.addWidget(self.canvas)
        self.setLayout(layout)

        self.plot_histogram(freq_data, title)

    def plot_histogram(self, freq_data, title):
        ax = self.figure.add_subplot(111)
        ax.clear()
        sorted_freq_data = dict(sorted(freq_data.items()))
        ax.bar(sorted_freq_data.keys(), sorted_freq_data.values(), color='skyblue')
        ax.set_title(title)
        ax.set_xlabel('Символы')
        ax.set_ylabel('Частота (%)')

        ax.set_xticks(range(len(sorted_freq_data)))
        ax.set_xticklabels(sorted_freq_data.keys(), rotation=0, ha='center')
        ax.tick_params(axis='x', which='both', labelsize=8)

        self.figure.tight_layout()

        self.canvas.draw()

class SubstitutionTableWindow(QtWidgets.QDialog):
    def __init__(self, language="ru", substitution_table=None, available_chars=None, freq_data=None):
        super().__init__()
        self.language = language
        self.setWindowTitle(f"Таблица замен ({'Русский' if language == 'ru' else 'Английский'})")
        self.resize(600, 700)
        self.setWindowModality(QtCore.Qt.ApplicationModal)
        
        # Создаем таблицу с 4 столбцами
        self.tableWidget = QtWidgets.QTableWidget(self)
        self.tableWidget.setColumnCount(4)
        self.tableWidget.setHorizontalHeaderLabels(["Исходный символ", "Замена", "Частота (%)", "Теор. вероятность (%)"])
        self.tableWidget.setSelectionBehavior(QtWidgets.QAbstractItemView.SelectRows)
        self.tableWidget.setSelectionMode(QtWidgets.QAbstractItemView.SingleSelection)
        self.tableWidget.setEditTriggers(QtWidgets.QAbstractItemView.NoEditTriggers)
        self.tableWidget.setAlternatingRowColors(True)
        self.tableWidget.setSortingEnabled(True)
        
        self.scrollArea = QtWidgets.QScrollArea(self)
        self.scrollArea.setWidgetResizable(True)
        self.scrollArea.setWidget(self.tableWidget)
        
        self.saveButton = QtWidgets.QPushButton("Сохранить", self)
        self.saveButton.clicked.connect(self.save_table)
        
        layout = QtWidgets.QVBoxLayout()
        layout.addWidget(self.scrollArea)
        layout.addWidget(self.saveButton)
        self.setLayout(layout)
        
        self.available_chars = available_chars if available_chars else set()
        self.freq_data = freq_data if freq_data else {}
        self.load_table(substitution_table)

    def load_table(self, substitution_table):
        # Выбираем соответствующий словарь вероятностей
        probabilities = RUSSIAN_PROBABILITIES if self.language == "ru" else ENGLISH_PROBABILITIES
        self.substitution_table = substitution_table if substitution_table else {char: char for char in probabilities.keys()}
        self.tableWidget.setRowCount(len(probabilities))
        
        for row, char in enumerate(sorted(probabilities.keys())):
            item_original = QtWidgets.QTableWidgetItem(char)
            item_original.setFlags(item_original.flags() & ~QtCore.Qt.ItemIsEditable)
            self.tableWidget.setItem(row, 0, item_original)
            
            line_edit = QtWidgets.QLineEdit(self.substitution_table[char])
            line_edit.setMaxLength(1)
            line_edit.textChanged.connect(lambda text, row=row: self.on_text_changed(text, row))
            self.tableWidget.setCellWidget(row, 1, line_edit)
            
            freq = self.freq_data.get(char, 0)
            item_freq = QtWidgets.QTableWidgetItem(f"{freq:.2f}" if freq else "-")
            item_freq.setFlags(item_freq.flags() & ~QtCore.Qt.ItemIsEditable)
            self.tableWidget.setItem(row, 2, item_freq)
            
            prob = probabilities[char] * 100  
            item_prob = QtWidgets.QTableWidgetItem(f"{prob:.2f}")
            item_prob.setFlags(item_prob.flags() & ~QtCore.Qt.ItemIsEditable)
            self.tableWidget.setItem(row, 3, item_prob)
        
        self.tableWidget.resizeColumnsToContents()

    def on_text_changed(self, text, row):
        original_char = self.tableWidget.item(row, 0).text()
        new_value = text

        if not new_value:
            return

        if self.language == "ru" and not all(c.isalpha() and c in RUSSIAN_PROBABILITIES for c in new_value):
            QtWidgets.QMessageBox.warning(self, "Ошибка", "В таблицу для русских символов можно вводить только русские символы!")
            line_edit = self.tableWidget.cellWidget(row, 1)
            line_edit.setText(self.substitution_table[original_char])
            return

        if self.language == "en" and not all(c.isalpha() and c in ENGLISH_PROBABILITIES for c in new_value):
            QtWidgets.QMessageBox.warning(self, "Ошибка", "В таблицу для английских символов можно вводить только английские символы!")
            line_edit = self.tableWidget.cellWidget(row, 1)
            line_edit.setText(self.substitution_table[original_char])
            return

        old_value = self.substitution_table[original_char]

        # Проверяем, есть ли дубликаты и обновляем соответствующие строки
        for r in range(self.tableWidget.rowCount()):
            if r == row:
                continue
            line_edit = self.tableWidget.cellWidget(r, 1)
            if line_edit.text() == new_value:
                # Найдено совпадение с другим символом
                other_char = self.tableWidget.item(r, 0).text()
                # Переназначаем старое значение для этой буквы
                line_edit.setText(old_value)
                self.substitution_table[other_char] = old_value
                break

        # Обновляем текущий символ
        self.substitution_table[original_char] = new_value

    def save_table(self):
        for row in range(self.tableWidget.rowCount()):
            original = self.tableWidget.item(row, 0).text()
            replacement = self.tableWidget.cellWidget(row, 1).text()
            self.substitution_table[original] = replacement
        QtWidgets.QMessageBox.information(self, "Успех", "Таблица замен сохранена!")
        self.accept()

class FrequencyAnalysisDialog(QtWidgets.QDialog):
    def __init__(self):
        super().__init__()
        self.setWindowIcon(QIcon("icon.ico"))
        self.setWindowTitle("Частотный анализ")
        self.resize(500, 400)
        self.setWindowModality(QtCore.Qt.ApplicationModal)
        self.textEdit = QtWidgets.QTextEdit(self)
        self.textEdit.setGeometry(20, 20, 460, 80)
        self.loadButton = QtWidgets.QPushButton("Загрузить из файла", self)
        self.loadButton.clicked.connect(self.load_from_file)
        self.clearButton = QtWidgets.QPushButton("Очистить", self)
        self.clearButton.clicked.connect(self.clear_text)
        self.plotButton = QtWidgets.QPushButton("Построить гистограмму", self)
        self.plotButton.clicked.connect(self.plot_histogram)
        self.ruSubstitutionButton = QtWidgets.QPushButton("Таблица замен (Русский)", self)
        self.ruSubstitutionButton.clicked.connect(lambda: self.open_substitution_table("ru"))
        self.enSubstitutionButton = QtWidgets.QPushButton("Таблица замен (Английский)", self)
        self.enSubstitutionButton.clicked.connect(lambda: self.open_substitution_table("en"))
        self.decryptButton = QtWidgets.QPushButton("Дешифровать текст", self)
        self.decryptButton.clicked.connect(self.decrypt_text)
        self.decryptedBrowser = QtWidgets.QTextBrowser(self)
        self.decryptedBrowser.setPlaceholderText("Дешифрованный текст будет отображён здесь")

        self.saveButton = QtWidgets.QPushButton("Сохранить", self)
        self.saveButton.clicked.connect(self.save_decrypted_text)
        
        layout = QtWidgets.QVBoxLayout()
        layout.addWidget(self.textEdit)
        buttonLayout = QtWidgets.QGridLayout()
        buttonLayout.addWidget(self.loadButton, 0, 0)
        buttonLayout.addWidget(self.clearButton, 0, 1)
        buttonLayout.addWidget(self.plotButton, 1, 0, 1, 2)
        buttonLayout.addWidget(self.ruSubstitutionButton, 2, 0)
        buttonLayout.addWidget(self.enSubstitutionButton, 2, 1)
        buttonLayout.addWidget(self.decryptButton, 3, 0, 1, 2)
        buttonLayout.addWidget(self.saveButton, 4, 0, 1, 2)
        layout.addLayout(buttonLayout)
        layout.addWidget(self.decryptedBrowser)
        self.setLayout(layout)
        self.freq_data_ru = None
        self.freq_data_en = None
        self.substitution_tables = {"ru": {}, "en": {}}
        self.available_chars = {"ru": set(), "en": set()}

    def load_from_file(self):
        options = QtWidgets.QFileDialog.Options()
        file_path, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Выбрать файл", "", "Текстовые файлы (*.txt *.docx);;Все файлы (*)", options=options)
        if file_path:
            _, ext = os.path.splitext(file_path)
            if ext.lower() == '.txt':
                with open(file_path, 'r', encoding='utf-8') as file:
                    self.textEdit.setText(file.read())
            elif ext.lower() == '.docx':
                doc = docx.Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                self.textEdit.setText(full_text)

    def clear_text(self):
        self.textEdit.clear()
        self.decryptedBrowser.clear()
        self.freq_data_ru = None
        self.freq_data_en = None
        self.substitution_tables = {"ru": {}, "en": {}}
        self.available_chars = {"ru": set(), "en": set()}

    def analyze_text(self):
        text = self.textEdit.toPlainText()

        text_lower = text.lower()
        text_filtered_ru = ''.join([char for char in text_lower if char.isalpha() and char in RUSSIAN_PROBABILITIES])
        text_filtered_en = ''.join([char for char in text_lower if char.isalpha() and char in ENGLISH_PROBABILITIES])

        if text_filtered_ru:
            freq_ru = Counter(text_filtered_ru)
            total_chars_ru = sum(freq_ru.values())
            self.freq_data_ru = {char: round(count / total_chars_ru * 100, 2) for char, count in freq_ru.items()}
            self.generate_substitution_table("ru")

        if text_filtered_en:
            freq_en = Counter(text_filtered_en)
            total_chars_en = sum(freq_en.values())
            self.freq_data_en = {char: round(count / total_chars_en * 100, 2) for char, count in freq_en.items()}
            self.generate_substitution_table("en")

    def generate_substitution_table(self, language):
        probabilities = RUSSIAN_PROBABILITIES if language == "ru" else ENGLISH_PROBABILITIES
        freq_data = self.freq_data_ru if language == "ru" else self.freq_data_en
        if not freq_data:
            return

        sorted_chars_by_freq = sorted(freq_data, key=freq_data.get, reverse=True)
        sorted_chars_by_prob = sorted(probabilities, key=probabilities.get, reverse=True)

        substitution_table = {}
        available_chars = set()

        for orig_char, prob_char in zip(sorted_chars_by_freq, sorted_chars_by_prob):
            if orig_char in probabilities:
                substitution_table[orig_char] = prob_char
                available_chars.add(prob_char)

        for orig_char in probabilities:
            if orig_char not in substitution_table:
                for prob_char in sorted_chars_by_prob:
                    if prob_char not in available_chars:
                        substitution_table[orig_char] = prob_char
                        available_chars.add(prob_char)
                        break

        self.substitution_tables[language] = substitution_table
        self.available_chars[language] = available_chars

    def plot_histogram(self):
        self.analyze_text()

        if not self.freq_data_ru and not self.freq_data_en:
            QtWidgets.QMessageBox.warning(self, "Ошибка", "Текст не содержит русских или английских символов для анализа")
            return

        self.hist_window = QtWidgets.QDialog(self)
        self.hist_window.setWindowTitle("Гистограммы")
        self.hist_window.setGeometry(150, 150, 800, 600)

        tab_widget = QtWidgets.QTabWidget(self.hist_window)

        if self.freq_data_ru:
            hist_widget_ru = HistogramWidget(self.freq_data_ru, "Гистограмма (Русский)")
            tab_widget.addTab(hist_widget_ru, "Русский")

        if self.freq_data_en:
            hist_widget_en = HistogramWidget(self.freq_data_en, "Гистограмма (Английский)")
            tab_widget.addTab(hist_widget_en, "Английский")

        layout = QtWidgets.QVBoxLayout()
        layout.addWidget(tab_widget)
        self.hist_window.setLayout(layout)

        self.hist_window.show()

    def open_substitution_table(self, language):
        substitution_table = self.substitution_tables.get(language)
        available_chars = self.available_chars.get(language)
        freq_data = self.freq_data_ru if language == "ru" else self.freq_data_en
        table_window = SubstitutionTableWindow(language, substitution_table, available_chars, freq_data)
        if table_window.exec_():
            self.substitution_tables[language] = table_window.substitution_table
            self.available_chars[language] = table_window.available_chars

    def decrypt_text(self):
        encrypted_text = self.textEdit.toPlainText()
        if not encrypted_text:
            QtWidgets.QMessageBox.warning(self, "Ошибка", "Введите текст для дешифровки")
            return

        if not self.substitution_tables["ru"] and not self.substitution_tables["en"]:
            QtWidgets.QMessageBox.warning(self, "Ошибка", "Сначала постройте гистрограмму!")
            return

        decrypted_text = []

        for char in encrypted_text:
            lower_char = char.lower()
            if lower_char in self.substitution_tables["ru"]:
                language = "ru"
            elif lower_char in self.substitution_tables["en"]:
                language = "en"
            else:
                decrypted_text.append(char)
                continue

            substitution_table = self.substitution_tables[language]
            if lower_char in substitution_table:
                if char.isupper():
                    decrypted_text.append(substitution_table[lower_char].upper())
                else:
                    decrypted_text.append(substitution_table[lower_char])
            else:
                decrypted_text.append(char)

        self.decryptedBrowser.setText(''.join(decrypted_text))

    def save_decrypted_text(self):
        decrypted_text = self.decryptedBrowser.toPlainText()
        if not decrypted_text:
            QtWidgets.QMessageBox.warning(self, "Ошибка", "Нет текста для сохранения")
            return
        options = QtWidgets.QFileDialog.Options()
        file_path, selected_filter = QtWidgets.QFileDialog.getSaveFileName(self, "Сохранить файл", "", "Текстовые файлы (*.txt);;DOCX файлы (*.docx);;Все файлы (*)", options=options)
        if file_path:
            try:
                _, ext = os.path.splitext(file_path)
                if ext.lower() == '.txt':
                    with open(file_path, 'w', encoding='utf-8') as file:
                        file.write(decrypted_text)
                elif ext.lower() == '.docx':
                    doc = docx.Document()
                    doc.add_paragraph(decrypted_text)
                    doc.save(file_path)
                QtWidgets.QMessageBox.information(self, "Успех", "Текст успешно сохранен!")
            except Exception as e:
                QtWidgets.QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить файл: {str(e)}")


if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    dialog = FrequencyAnalysisDialog()
    dialog.show()
    sys.exit(app.exec_())
